# -*- coding: utf-8 -*-
import os
import zipfile
from datetime import datetime

import geopandas as gpd
import matplotlib.image as mpimg
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from matplotlib import font_manager
from matplotlib.colors import BoundaryNorm, ListedColormap
from matplotlib.lines import Line2D
from matplotlib.patches import Patch, Rectangle
from shapely.geometry import Point, Polygon

try:
    import contextily as ctx
except Exception:
    ctx = None

try:
    import rasterio
    from rasterio.enums import Resampling
    from rasterio.features import geometry_mask
    from rasterio.fill import fillnodata
    from rasterio.mask import mask as raster_mask
    from rasterio.vrt import WarpedVRT
except Exception:
    rasterio = None
    geometry_mask = None
    fillnodata = None
    raster_mask = None
    WarpedVRT = None
    Resampling = None


CRS_EPSG = "EPSG:32644"
DEFAULT_DEM_PATH = "/home/ubuntu/gfcnepal/dem/Nepal.tif"
DEM_BUFFER_M = 10
CRS_OPTIONS = {
    "EPSG:32644": "CRS Name: WGS 84 / UTM zone 44N\nEPSG:32644",
    "EPSG:32645": "CRS Name: WGS 84 / UTM zone 45N\nEPSG:32645",
}

NORTH_ARROW_IMAGE = "/home/ubuntu/gfcnepal/static/north_arrow.png"

ALLOWED_MAPS = [
    "boundary_map",
    "compartment_map",
    "slope_map",
    "elevation_map",
    "aspect_map",
    "plot_map",
    "google_earth_map",
    "silviculture_map",
    "mother_tree_map",
]

MAP_TITLES = {
    "boundary_map": "Boundary and Survey Points Map",
    "compartment_map": "Compartment Map",
    "slope_map": "Slope Map",
    "elevation_map": "Elevation Map",
    "aspect_map": "Aspect Map",
    "plot_map": "Inventory Plot Map",
    "google_earth_map": "Satellite View Map",
    "silviculture_map": "Silviculture Treatment Map",
    "mother_tree_map": "Mother Tree Plot Map",
}

BLOCK_COLORS = [
    "#ff0000", "#00ff00", "#008b8b", "#808080", "#ff00ff",
    "#00ffff", "#8b0000", "#008000", "#c0c0c0", "#808000",
]
SLOPE_COLORS = ["#2fb344", "#f2d64b", "#8f9699", "#d9822b", "#b91c1c"]
ELEVATION_COLORS = ["#d9f0a3", "#78c679", "#41b6c4", "#2b8cbe", "#253494"]
ASPECT_COLORS = ["#2166ac", "#67a9cf", "#1a9850", "#a6d96a", "#fdae61", "#f46d43", "#d73027", "#762a83"]
SLOPE_BINS = [0, 8.5, 19, 31, 45, np.inf]
SLOPE_LABELS = ["0-8.5 deg", "8.5-19 deg", "19-31 deg", "31-45 deg", ">45 deg"]
ASPECT_LABELS = ["N", "NE", "E", "SE", "S", "SW", "W", "NW"]

DEVANAGARI_FONT = None
DEVANAGARI_FONT_BOLD = None


def set_font():
    global DEVANAGARI_FONT, DEVANAGARI_FONT_BOLD
    candidates = [
        r"C:\Windows\Fonts\Nirmala.ttf",
        r"C:\Windows\Fonts\Mangal.ttf",
        r"C:\Windows\Fonts\kokila.ttf",
        r"C:\Windows\Fonts\aparaj.ttf",
        r"C:\Windows\Fonts\arialuni.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    bold_candidates = [
        r"C:\Windows\Fonts\NirmalaB.ttf",
        r"C:\Windows\Fonts\Mangal.ttf",
        r"C:\Windows\Fonts\kokilab.ttf",
        r"C:\Windows\Fonts\aparajb.ttf",
        r"C:\Windows\Fonts\arialuni.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]
    font_path = next((p for p in candidates if os.path.exists(p)), None)
    bold_path = next((p for p in bold_candidates if os.path.exists(p)), font_path)
    if font_path:
        font_manager.fontManager.addfont(font_path)
    if bold_path:
        font_manager.fontManager.addfont(bold_path)
    if font_path:
        DEVANAGARI_FONT = font_manager.FontProperties(fname=font_path)
        DEVANAGARI_FONT_BOLD = font_manager.FontProperties(fname=bold_path, weight="bold")
        plt.rcParams["font.family"] = DEVANAGARI_FONT.get_name()
    else:
        DEVANAGARI_FONT = font_manager.FontProperties(family="DejaVu Sans")
        DEVANAGARI_FONT_BOLD = font_manager.FontProperties(family="DejaVu Sans", weight="bold")
        plt.rcParams["font.family"] = "DejaVu Sans"
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["pdf.fonttype"] = 42
    plt.rcParams["ps.fonttype"] = 42


def nep_font(bold=False):
    return DEVANAGARI_FONT_BOLD if bold else DEVANAGARI_FONT


def get_crs_text(crs_epsg):
    crs_epsg = normalize_crs(crs_epsg)
    return CRS_OPTIONS.get(crs_epsg, CRS_OPTIONS[CRS_EPSG])


def normalize_crs(crs_epsg):
    text = str(crs_epsg or "").strip().upper().replace(" ", "")
    if text in ["45", "45N", "ZONE45", "ZONE45N", "32645", "EPSG32645", "EPSG:32645"]:
        return "EPSG:32645"
    if text in ["44", "44N", "ZONE44", "ZONE44N", "32644", "EPSG32644", "EPSG:32644"]:
        return "EPSG:32644"
    if text == "EPSG:":
        return "EPSG:32645"
    return text or CRS_EPSG


def is_shapefile_input(file_path):
    return os.path.splitext(file_path)[1].lower() in [".shp", ".zip", ".gpkg", ".geojson", ".json"]


def resolve_boundary_input(input_path):
    if not os.path.isdir(input_path):
        return input_path

    supported_exts = [".shp", ".gpkg", ".zip", ".geojson", ".json", ".xlsx", ".xls", ".csv"]
    files = []
    for name in os.listdir(input_path):
        full_path = os.path.join(input_path, name)
        if os.path.isfile(full_path) and os.path.splitext(name)[1].lower() in supported_exts:
            files.append(full_path)

    if not files:
        raise FileNotFoundError(
            "No supported boundary file found in folder: {}. Use .shp, .zip, .gpkg, .geojson, .xlsx, .xls, or .csv.".format(input_path)
        )

    def priority(path):
        name = os.path.splitext(os.path.basename(path))[0].lower()
        ext = os.path.splitext(path)[1].lower()
        boundary_rank = 0 if name == "boundary" else 1
        ext_rank = {
            ".shp": 0,
            ".gpkg": 1,
            ".zip": 2,
            ".geojson": 3,
            ".json": 4,
            ".xlsx": 5,
            ".xls": 6,
            ".csv": 7,
        }.get(ext, 99)
        return boundary_rank, ext_rank, os.path.basename(path).lower()

    return sorted(files, key=priority)[0]


def read_shapefile_blocks(file_path, crs_epsg=CRS_EPSG):
    crs_epsg = normalize_crs(crs_epsg)
    if file_path.lower().endswith(".zip"):
        with zipfile.ZipFile(file_path, "r") as zf:
            if not any(n.lower().endswith(".shp") for n in zf.namelist()):
                raise ValueError("No .shp file was found inside the ZIP file.")
        gdf = gpd.read_file("zip://{}".format(file_path))
    elif file_path.lower().endswith(".gpkg"):
        gdf = gpd.read_file(file_path)
    else:
        gdf = gpd.read_file(file_path)
    if gdf.empty:
        raise ValueError("The boundary file is empty.")
    gdf = gdf.set_crs(crs_epsg) if gdf.crs is None else gdf.to_crs(crs_epsg)
    gdf = gdf[gdf.geometry.notna()].copy()
    gdf = gdf[gdf.geometry.geom_type.isin(["Polygon", "MultiPolygon"])].copy()
    if gdf.empty:
        raise ValueError("No polygon geometry was found in the shapefile.")
    block_col = next((c for c in ["Name", "name", "Block name", "block", "Block", "BLOCK_NAME", "AWADI", "Awadi", "awadi"] if c in gdf.columns), None)
    if len(gdf) == 1:
        gdf["block"] = "Forest Area"
    elif block_col:
        gdf["block"] = gdf[block_col].fillna("").astype(str).str.strip()
        empty_labels = gdf["block"] == ""
        gdf.loc[empty_labels, "block"] = ["Block {}".format(i + 1) for i in range(int(empty_labels.sum()))]
    else:
        gdf["block"] = ["Block {}".format(i + 1) for i in range(len(gdf))]
    return gdf[["block", "geometry"]].reset_index(drop=True)


def read_survey_table(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    df = pd.read_csv(file_path) if ext == ".csv" else pd.read_excel(file_path)
    df.columns = [str(c).strip() for c in df.columns]
    required = ["SN", "Block name", "VertexIndex", "X", "Y"]
    missing = [c for c in required if c not in df.columns]
    if missing:
        raise ValueError("Missing columns: {}".format(", ".join(missing)))
    df = df.dropna(subset=["Block name", "VertexIndex", "X", "Y"]).copy()
    df["VertexIndex"] = pd.to_numeric(df["VertexIndex"], errors="coerce")
    df["X"] = pd.to_numeric(df["X"], errors="coerce")
    df["Y"] = pd.to_numeric(df["Y"], errors="coerce")
    return df.dropna(subset=["VertexIndex", "X", "Y"]).sort_values(["Block name", "VertexIndex"])


def build_layers(df, crs_epsg=CRS_EPSG):
    crs_epsg = normalize_crs(crs_epsg)
    polygon_rows, point_rows = [], []
    for block_name, group in df.groupby("Block name", sort=False):
        coords = list(zip(group["X"], group["Y"]))
        if len(coords) < 3:
            continue
        if coords[0] != coords[-1]:
            coords.append(coords[0])
        poly = Polygon(coords)
        if not poly.is_valid:
            poly = poly.buffer(0)
        polygon_rows.append({"block": str(block_name), "geometry": poly})
        for _, row in group.iterrows():
            point_rows.append({"block": str(block_name), "vertex": int(row["VertexIndex"]), "geometry": Point(float(row["X"]), float(row["Y"]))})
    blocks = gpd.GeoDataFrame(polygon_rows, crs=crs_epsg)
    if len(blocks) == 1:
        blocks.loc[blocks.index[0], "block"] = "Forest Area"
    return blocks, gpd.GeoDataFrame(point_rows, crs=crs_epsg)


def build_layers_from_input(input_file, crs_epsg=CRS_EPSG):
    crs_epsg = normalize_crs(crs_epsg)
    input_file = resolve_boundary_input(input_file)
    if is_shapefile_input(input_file):
        blocks = read_shapefile_blocks(input_file, crs_epsg=crs_epsg)
        points = gpd.GeoDataFrame(columns=["block", "vertex", "geometry"], crs=crs_epsg)
        return blocks, points
    return build_layers(read_survey_table(input_file), crs_epsg=crs_epsg)


def total_polygon_area_ha(blocks):
    return float(blocks.geometry.area.sum() / 10000.0)


def polygon_areas_ha(blocks):
    return [float(geom.area / 10000.0) for geom in blocks.geometry]


def polygon_area_labels(blocks):
    if len(blocks) == 1:
        return ["Forest Area"]
    return [str(row["block"]) for _, row in blocks.iterrows()]


def polygon_area_title(blocks):
    return "Area" if len(blocks) == 1 else "Block Area"


def polygon_area_col1(blocks):
    return "Item" if len(blocks) == 1 else "Block"


def setup_axes(ax, blocks):
    minx, miny, maxx, maxy = blocks.total_bounds
    width, height = maxx - minx, maxy - miny
    pad = max(width, height) * 0.10
    ax.set_xlim(minx - pad, maxx + pad)
    ax.set_ylim(miny - pad, maxy + pad)
    ax.set_aspect("equal", adjustable="box")
    ax.set_xticks([minx, (minx + maxx) / 2, maxx])
    ax.set_yticks([miny, miny + height / 3, miny + 2 * height / 3, maxy])
    ax.ticklabel_format(style="plain", useOffset=False)
    ax.grid(True, color="#8a8a8a", linewidth=0.35, alpha=0.65)
    ax.tick_params(
        axis="x",
        which="both",
        top=True,
        bottom=True,
        labeltop=True,
        labelbottom=True,
        labelsize=5.5,
        pad=1,
    )
    ax.tick_params(
        axis="y",
        which="both",
        left=True,
        right=True,
        labelleft=True,
        labelright=True,
        labelsize=5.5,
        pad=1,
        labelrotation=90,
    )
    for spine in ax.spines.values():
        spine.set_visible(False)


def add_title_band(fig, map_key, forest_name, address):
    """Pillow ले Nepali title (forest name, address) र English map title render गर्छ"""
    from PIL import Image, ImageDraw, ImageFont
    import numpy as np

    map_title = MAP_TITLES.get(map_key, map_key)
    line1 = "{}, {}".format(forest_name, address)   # Nepali - Noto Devanagari
    line2 = map_title                                 # English - Times New Roman

    # Font paths
    nep_bold   = "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Bold.ttf"
    nep_reg    = "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf"
    eng_bold   = "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf"
    eng_reg    = "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf"

    # Fallbacks
    if not os.path.exists(nep_bold):   nep_bold  = nep_reg
    if not os.path.exists(nep_reg):    nep_reg   = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    if not os.path.exists(eng_bold):   eng_bold  = nep_bold
    if not os.path.exists(eng_reg):    eng_reg   = nep_reg

    fig_width_inch, fig_height_inch = fig.get_size_inches()
    dpi = fig.dpi
    fig_w_px = int(fig_width_inch * dpi)

    # Font sizes proportional to figure width
    f1 = max(24, int(fig_w_px * 0.032))   # Nepali line (forest name, address)
    f2 = max(20, int(fig_w_px * 0.024))   # English line (map title)
    band_h = f1 + f2 + 50

    img  = Image.new("RGBA", (fig_w_px, band_h), (255, 255, 255, 255))
    draw = ImageDraw.Draw(img)

    try:
        font1 = ImageFont.truetype(nep_bold, f1)
    except Exception:
        font1 = ImageFont.load_default()
    try:
        font2 = ImageFont.truetype(eng_bold, f2)
    except Exception:
        font2 = ImageFont.load_default()

    # Center line1 (Nepali)
    bb1 = draw.textbbox((0, 0), line1, font=font1)
    w1  = bb1[2] - bb1[0]
    draw.text(((fig_w_px - w1) // 2, 10), line1, font=font1, fill=(0, 0, 0, 255))

    # Center line2 (English)
    bb2 = draw.textbbox((0, 0), line2, font=font2)
    w2  = bb2[2] - bb2[0]
    draw.text(((fig_w_px - w2) // 2, f1 + 20), line2, font=font2, fill=(50, 50, 50, 255))

    # Place band at top of figure
    frac = band_h / (fig_height_inch * dpi)
    ax_title = fig.add_axes([0.0, 1.0 - frac, 1.0, frac])
    ax_title.axis("off")
    ax_title.imshow(np.array(img), aspect="auto", interpolation="bilinear")


def add_north_arrow(ax):
    if os.path.exists(NORTH_ARROW_IMAGE):
        image = mpimg.imread(NORTH_ARROW_IMAGE)
        inset = ax.inset_axes([0.905, 0.815, 0.075, 0.075], transform=ax.transAxes)
        inset.imshow(image)
        inset.axis("off")
        return
    ax.annotate("N", xy=(0.945, 0.90), xytext=(0.945, 0.82), xycoords="axes fraction", ha="center", va="center", fontsize=8, fontweight="bold", arrowprops=dict(arrowstyle="-|>", mutation_scale=12, lw=1.8, color="black"), zorder=20)


def add_scale_bar(fig, length_m=500):
    ax = fig.add_axes([0.650, 0.080, 0.220, 0.030])
    ax.axis("off")
    ax.add_patch(Rectangle((0.00, 0.38), 0.50, 0.22, facecolor="black", edgecolor="black"))
    ax.add_patch(Rectangle((0.50, 0.38), 0.50, 0.22, facecolor="white", edgecolor="black", linewidth=0.8))
    ax.text(0.00, 0.10, "0", ha="center", va="top", fontsize=7)
    ax.text(0.50, 0.10, str(int(length_m / 2)), ha="center", va="top", fontsize=7)
    ax.text(1.00, 0.10, "{} m".format(length_m), ha="center", va="top", fontsize=7)


def add_crs_text(fig, crs_epsg=CRS_EPSG):
    from matplotlib.font_manager import FontProperties
    eng_font = FontProperties(fname="/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf") if os.path.exists("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf") else FontProperties()
    crs_epsg = normalize_crs(crs_epsg)
    fig.text(0.650, 0.120, get_crs_text(crs_epsg), ha="left", va="bottom", fontsize=6.3, fontproperties=eng_font)


def add_area_table(fig, labels, areas_ha, title="Area Table", col1="Class", col2="Area", position=None):
    rows, total = [], 0.0
    for label, area in zip(labels, areas_ha):
        area = float(area)
        if area <= 0:
            continue
        total += area
        rows.append([str(label), "{:.2f} ha".format(area)])
    rows.append(["Total", "{:.2f} ha".format(total)])
    ax = fig.add_axes(position or [0.370, 0.070, 0.245, 0.115])
    ax.axis("off")
    ax.text(0.0, 1.02, title, ha="left", va="bottom", fontsize=7, fontweight="bold")
    table = ax.table(cellText=rows, colLabels=[col1, col2], loc="upper left", cellLoc="left", colLoc="left", bbox=[0, 0, 1, 0.98])
    table.auto_set_font_size(False)
    table.set_fontsize(5.6)
    for (row, _), cell in table.get_celld().items():
        cell.set_linewidth(0.35)
        cell.set_edgecolor("#444444")
        if row == 0:
            cell.set_facecolor("#eeeeee")
            cell.set_text_props(weight="bold")


def add_boundary_legend(fig):
    from matplotlib.font_manager import FontProperties
    eng_font = FontProperties(fname="/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf") if os.path.exists("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf") else FontProperties()
    eng_bold = FontProperties(fname="/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf") if os.path.exists("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf") else FontProperties()
    handles = [
        Line2D([0], [0], marker="o", color="black", linestyle="", markersize=4, label="Survey Points"),
        Patch(facecolor="white", edgecolor="black", label="Forest Boundary"),
    ]
    fig.legend(handles=handles, title="Legend", loc="lower left", bbox_to_anchor=(0.105, 0.095), frameon=False, prop=eng_font, title_fontproperties=eng_bold, fontsize=7)


def add_block_legend(fig, blocks, title="Legend", colors=None, labels=None, position=None, ncol=2):
    from matplotlib.font_manager import FontProperties
    eng_font = FontProperties(fname="/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf") if os.path.exists("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf") else FontProperties()
    eng_bold = FontProperties(fname="/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf") if os.path.exists("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf") else FontProperties()
    nep_font_prop = nep_font()
    colors = colors or BLOCK_COLORS
    labels = labels or [str(row["block"]) for _, row in blocks.iterrows()]
    handles = [Patch(facecolor=colors[i % len(colors)], edgecolor="black", label=labels[i]) for i in range(min(len(colors), len(labels)))]
    fig.legend(handles=handles, title=title, loc="lower left", bbox_to_anchor=position or (0.105, 0.090), frameon=False, ncol=ncol, prop=nep_font_prop, title_fontproperties=eng_bold, fontsize=7, borderaxespad=0, columnspacing=0.75, handlelength=1.0)


def class_areas_ha(class_data, class_count, pixel_x, pixel_y, target_total_ha=None):
    pixel_area_ha = abs(pixel_x * pixel_y) / 10000.0
    areas = [float(np.nansum(class_data == idx) * pixel_area_ha) for idx in range(class_count)]
    if target_total_ha and sum(areas) > 0:
        factor = float(target_total_ha) / float(sum(areas))
        areas = [area * factor for area in areas]
    return areas


def raster_extent(transform, shape):
    rows, cols = shape
    left, top = transform.c, transform.f
    right, bottom = left + cols * transform.a, top + rows * transform.e
    return [min(left, right), max(left, right), min(bottom, top), max(bottom, top)]


def fill_dem_for_gradient(data):
    valid_mask = np.isfinite(data)
    if not np.any(valid_mask):
        return data
    if np.all(valid_mask):
        return data

    if fillnodata is not None:
        filled = fillnodata(
            data.astype("float32"),
            mask=valid_mask.astype("uint8"),
            max_search_distance=50,
            smoothing_iterations=0,
        ).astype("float64")
        filled[~np.isfinite(filled)] = np.nanmean(data)
        return filled

    filled = data.copy()
    fill_value = float(np.nanmean(data))
    filled[~valid_mask] = fill_value
    return filled


def load_dem_layer(blocks, dem_path, layer_type):
    if not dem_path or layer_type not in ["elevation_map", "slope_map", "aspect_map"]:
        return None
    if rasterio is None:
        raise ImportError("DEM analysis needs rasterio. Please install rasterio.")
    if geometry_mask is None:
        raise ImportError("DEM analysis needs rasterio.features.geometry_mask.")
    if not os.path.exists(dem_path):
        raise FileNotFoundError("DEM file was not found: {}".format(dem_path))

    dem_buffer = blocks.copy()
    dem_buffer["geometry"] = dem_buffer.geometry.buffer(DEM_BUFFER_M)
    dem_buffer = dem_buffer[dem_buffer.geometry.notna() & ~dem_buffer.geometry.is_empty].copy()
    if dem_buffer.empty:
        dem_buffer = blocks.copy()

    with rasterio.open(dem_path) as src:
        if not src.crs:
            raise ValueError("The DEM file has no CRS. Please use a GeoTIFF with CRS metadata.")
        with WarpedVRT(src, crs=blocks.crs, resampling=Resampling.bilinear, nodata=src.nodata) as vrt:
            arr, transform = raster_mask(vrt, dem_buffer.geometry, crop=True, filled=True, nodata=vrt.nodata)
            data = arr[0].astype("float64")
            if vrt.nodata is not None:
                data[data == vrt.nodata] = np.nan
            data[~np.isfinite(data)] = np.nan
            pixel_x, pixel_y = abs(transform.a) or 1, abs(transform.e) or abs(transform.a) or 1
            inside_polygon = geometry_mask(
                blocks.geometry,
                out_shape=data.shape,
                transform=transform,
                invert=True,
                all_touched=True,
            )
    if np.all(np.isnan(data)):
        return None

    polygon_data = data.copy()
    polygon_data[~inside_polygon] = np.nan
    total_ha = total_polygon_area_ha(blocks)
    if layer_type == "elevation_map":
        valid = polygon_data[np.isfinite(polygon_data)]
        breaks = np.linspace(float(np.nanmin(valid)), float(np.nanmax(valid)), 6)
        if np.unique(breaks).size < 2:
            breaks = np.array([float(np.nanmin(valid)), float(np.nanmax(valid)) + 1])
        class_data = np.digitize(polygon_data, breaks[1:-1], right=True).astype("float64")
        class_data[~np.isfinite(polygon_data)] = np.nan
        labels = ["{:.0f}-{:.0f} m".format(breaks[i], breaks[i + 1]) for i in range(len(breaks) - 1)]
        return {"data": class_data, "colors": ELEVATION_COLORS, "labels": labels, "areas_ha": class_areas_ha(class_data, len(labels), pixel_x, pixel_y, total_ha), "extent": raster_extent(transform, data.shape), "title": "Elevation Area", "col1": "Elevation"}
    gradient_data = fill_dem_for_gradient(data)
    grad_y, grad_x = np.gradient(gradient_data, pixel_y, pixel_x)
    if layer_type == "slope_map":
        slope = np.degrees(np.arctan(np.sqrt((grad_x ** 2) + (grad_y ** 2))))
        class_data = np.digitize(slope, SLOPE_BINS[1:-1], right=True).astype("float64")
        class_data[(~np.isfinite(slope)) | (~inside_polygon)] = np.nan
        return {"data": class_data, "colors": SLOPE_COLORS, "labels": SLOPE_LABELS, "areas_ha": class_areas_ha(class_data, len(SLOPE_LABELS), pixel_x, pixel_y, total_ha), "extent": raster_extent(transform, data.shape), "title": "Slope Area", "col1": "Slope"}
    aspect = (np.degrees(np.arctan2(-grad_x, grad_y)) + 360) % 360
    class_data = np.zeros_like(aspect, dtype="float64")
    class_data[(aspect >= 337.5) | (aspect < 22.5)] = 0
    class_data[(aspect >= 22.5) & (aspect < 67.5)] = 1
    class_data[(aspect >= 67.5) & (aspect < 112.5)] = 2
    class_data[(aspect >= 112.5) & (aspect < 157.5)] = 3
    class_data[(aspect >= 157.5) & (aspect < 202.5)] = 4
    class_data[(aspect >= 202.5) & (aspect < 247.5)] = 5
    class_data[(aspect >= 247.5) & (aspect < 292.5)] = 6
    class_data[(aspect >= 292.5) & (aspect < 337.5)] = 7
    class_data[(~np.isfinite(aspect)) | (~inside_polygon)] = np.nan
    return {"data": class_data, "colors": ASPECT_COLORS, "labels": ASPECT_LABELS, "areas_ha": class_areas_ha(class_data, len(ASPECT_LABELS), pixel_x, pixel_y, total_ha), "extent": raster_extent(transform, data.shape), "title": "Aspect Area", "col1": "Aspect"}


def draw_dem_layer(ax, fig, blocks, dem_layer):
    colors, data = dem_layer["colors"], dem_layer["data"]
    max_class = int(np.nanmax(data)) if np.any(np.isfinite(data)) else len(colors) - 1
    cmap = ListedColormap(colors[:max_class + 1])
    norm = BoundaryNorm(np.arange(-0.5, len(colors) + 0.5, 1), cmap.N)
    ax.imshow(
        data,
        extent=dem_layer["extent"],
        origin="upper",
        cmap=cmap,
        norm=norm,
        alpha=0.96,
        interpolation="nearest",
        resample=False,
        zorder=1,
    )
    blocks.boundary.plot(ax=ax, color="#111111", linewidth=1.05, zorder=3)
    for _, row in blocks.iterrows():
        center = row.geometry.representative_point()
        ax.text(center.x, center.y, row["block"], color="#111111", fontsize=7.5, ha="center", va="center", zorder=4, fontproperties=nep_font())
    add_block_legend(fig, blocks, title="Legend", colors=colors, labels=dem_layer["labels"], ncol=2)
    add_area_table(fig, dem_layer["labels"], dem_layer["areas_ha"], title=dem_layer["title"], col1=dem_layer["col1"], col2="Hectare")


def draw_boundary_map(ax, fig, blocks, points):
    blocks.dissolve().boundary.plot(ax=ax, color="black", linewidth=1.4)
    blocks.boundary.plot(ax=ax, color="black", linewidth=0.9)
    if points is not None and not points.empty:
        points.plot(ax=ax, color="black", markersize=9, zorder=5)
        for _, row in points.iterrows():
            ax.annotate(str(row["vertex"]), xy=(row.geometry.x, row.geometry.y), xytext=(2, 2), textcoords="offset points", fontsize=5.5, color="black", zorder=6)
    for _, row in blocks.iterrows():
        center = row.geometry.representative_point()
        ax.text(center.x, center.y, row["block"], color="red", fontsize=9, ha="center", va="center", zorder=7, fontproperties=nep_font())
    add_boundary_legend(fig)
    add_area_table(
        fig,
        polygon_area_labels(blocks),
        polygon_areas_ha(blocks),
        title=polygon_area_title(blocks),
        col1=polygon_area_col1(blocks),
        col2="Hectare",
        position=[0.370, 0.075, 0.245, 0.110 if len(blocks) > 1 else 0.060],
    )


def draw_compartment_map(ax, fig, blocks):
    labels = []
    for i, row in blocks.iterrows():
        color = BLOCK_COLORS[i % len(BLOCK_COLORS)]
        labels.append(str(row["block"]))
        gpd.GeoSeries([row.geometry], crs=blocks.crs).plot(ax=ax, facecolor=color, edgecolor="black", linewidth=0.9, alpha=0.95)
        center = row.geometry.representative_point()
        ax.text(center.x, center.y, row["block"], color="black", fontsize=9, ha="center", va="center", fontproperties=nep_font())
    add_block_legend(fig, blocks, title="Block", colors=BLOCK_COLORS, labels=labels, ncol=2)
    add_area_table(fig, labels, polygon_areas_ha(blocks), title="Block Area", col1="Block", col2="Hectare")


def add_google_basemap(ax, blocks):
    if ctx is None:
        return False
    try:
        web_blocks = blocks.to_crs(epsg=3857)
        minx, miny, maxx, maxy = web_blocks.total_bounds
        pad = max(maxx - minx, maxy - miny) * 0.08
        ax.set_xlim(minx - pad, maxx + pad)
        ax.set_ylim(miny - pad, maxy + pad)
        ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, zoom="auto", attribution=False)
        web_blocks.boundary.plot(ax=ax, color="yellow", linewidth=1.2, zorder=5)
        for _, row in web_blocks.iterrows():
            center = row.geometry.representative_point()
            ax.text(center.x, center.y, row["block"], color="white", fontsize=8, ha="center", va="center", fontweight="bold", zorder=6)
        ax.set_axis_off()
        return True
    except Exception:
        return False


def draw_thematic_map(ax, fig, blocks, map_key, dem_path=None):
    dem_layer = load_dem_layer(blocks, dem_path, map_key)
    if dem_layer:
        draw_dem_layer(ax, fig, blocks, dem_layer)
        return
    if map_key == "google_earth_map" and add_google_basemap(ax, blocks):
        add_area_table(fig, polygon_area_labels(blocks), polygon_areas_ha(blocks), title=polygon_area_title(blocks), col1=polygon_area_col1(blocks), col2="Hectare")
        return
    palettes = {
        "slope_map": SLOPE_COLORS,
        "elevation_map": ELEVATION_COLORS,
        "aspect_map": ASPECT_COLORS,
        "plot_map": ["#f7f7f7", "#d9ead3"],
        "google_earth_map": ["#355e3b", "#4f7942", "#6b8e23", "#2e8b57"],
        "silviculture_map": ["#d9f0a3", "#78c679", "#238443", "#005a32"],
        "mother_tree_map": ["#f7fcb9", "#addd8e", "#31a354"],
    }
    class_labels = {
        "slope_map": ["0-8.5", "8.5-19", "19-31", "31-45", ">45"],
        "elevation_map": ["Low", "Low-Mid", "Mid", "Mid-High", "High"],
        "aspect_map": ["North", "North-East", "East", "South-East", "South"],
        "plot_map": ["Plot grid", "Forest area"],
        "google_earth_map": ["Google imagery layer", "Forest overlay", "Block boundary", "Label"],
        "silviculture_map": ["AAH", "TI", "RP", "Protection"],
        "mother_tree_map": ["Mother tree", "Spacing zone", "Forest area"],
    }
    colors = palettes.get(map_key, BLOCK_COLORS)
    for i, row in blocks.iterrows():
        gpd.GeoSeries([row.geometry], crs=blocks.crs).plot(ax=ax, facecolor=colors[i % len(colors)], edgecolor="black", linewidth=0.9, alpha=0.92)
        center = row.geometry.representative_point()
        ax.text(center.x, center.y, row["block"], color="red", fontsize=8, ha="center", va="center", fontproperties=nep_font())
    add_block_legend(fig, blocks, title="Legend", colors=colors, labels=class_labels.get(map_key))
    add_area_table(fig, polygon_area_labels(blocks), polygon_areas_ha(blocks), title=polygon_area_title(blocks), col1=polygon_area_col1(blocks), col2="Hectare")


def make_portrait_doc(docx_path, image_paths):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for margin in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, margin, Inches(0.35))
    for index, image_path in enumerate(image_paths):
        if index:
            doc.add_page_break()
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.add_run().add_picture(image_path, width=Inches(7.75))
    doc.save(docx_path)


def create_map_docx(docx_path, png_path):
    make_portrait_doc(docx_path, [png_path])


def create_all_maps_docx(docx_path, generated_maps, output_dir):
    make_portrait_doc(docx_path, [os.path.join(output_dir, item["png"]) for item in generated_maps])


def generate_map_from_layers(blocks, points, map_key, forest_name, address, output_png, output_pdf=None, crs_epsg=CRS_EPSG, dem_path=DEFAULT_DEM_PATH):
    crs_epsg = normalize_crs(crs_epsg)
    if map_key not in ALLOWED_MAPS:
        raise ValueError("Unknown map type: {}".format(map_key))
    if crs_epsg not in CRS_OPTIONS:
        raise ValueError("Unsupported CRS: {}. Use EPSG:32644 or EPSG:32645.".format(crs_epsg))
    if blocks.empty:
        raise ValueError("No valid block polygon found.")
    set_font()
    fig = plt.figure(figsize=(8.27, 11.69), dpi=400)
    ax = fig.add_axes([0.07, 0.235, 0.86, 0.745])
    add_title_band(fig, map_key, forest_name, address)
    setup_axes(ax, blocks)
    if map_key == "boundary_map":
        draw_boundary_map(ax, fig, blocks, points)
    elif map_key == "compartment_map":
        draw_compartment_map(ax, fig, blocks)
    else:
        draw_thematic_map(ax, fig, blocks, map_key, dem_path=dem_path)
    add_north_arrow(ax)
    add_crs_text(fig, crs_epsg=crs_epsg)
    add_scale_bar(fig)
    fig.savefig(output_png, dpi=400, facecolor="white")
    if output_pdf:
        fig.savefig(output_pdf, dpi=400, facecolor="white")
    plt.close(fig)


def generate_map(input_file, map_key, forest_name, address, output_png, output_pdf=None, crs_epsg=CRS_EPSG, dem_path=DEFAULT_DEM_PATH):
    crs_epsg = normalize_crs(crs_epsg)
    blocks, points = build_layers_from_input(input_file, crs_epsg=crs_epsg)
    generate_map_from_layers(blocks, points, map_key, forest_name, address, output_png, output_pdf, crs_epsg, dem_path)


def generate_selected_maps(input_file, selected_maps, forest_name, address, output_dir, crs_epsg=CRS_EPSG, dem_path=DEFAULT_DEM_PATH):
    crs_epsg = normalize_crs(crs_epsg)
    os.makedirs(output_dir, exist_ok=True)
    selected_maps = [m for m in selected_maps if m in ALLOWED_MAPS]
    if not selected_maps:
        raise ValueError("Please select at least one map.")
    blocks, points = build_layers_from_input(input_file, crs_epsg=crs_epsg)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    generated = []
    for index_no, map_key in enumerate(selected_maps, start=1):
        png_name = "map{}_{}_{}.png".format(index_no, map_key, timestamp)
        pdf_name = "map{}_{}_{}.pdf".format(index_no, map_key, timestamp)
        docx_name = "map{}_{}_{}.docx".format(index_no, map_key, timestamp)
        png_path = os.path.join(output_dir, png_name)
        pdf_path = os.path.join(output_dir, pdf_name)
        docx_path = os.path.join(output_dir, docx_name)
        generate_map_from_layers(blocks, points, map_key, forest_name, address, png_path, pdf_path, crs_epsg=crs_epsg, dem_path=dem_path)
        create_map_docx(docx_path, png_path)
        generated.append({"key": map_key, "title": MAP_TITLES[map_key], "png": png_name, "pdf": pdf_name, "docx": docx_name})
    all_docx_name = "all_maps_{}.docx".format(timestamp)
    create_all_maps_docx(os.path.join(output_dir, all_docx_name), generated, output_dir)
    for item in generated:
        item["all_docx"] = all_docx_name
    return generated

if __name__ == "__main__":
    generate_selected_maps(
        input_file=r"/home/ubuntu/gfcnepal/BoundaryCF.xlsx",
        selected_maps=[
            "boundary_map",
            "elevation_map",
            "slope_map",
            "aspect_map",
            "google_earth_map",
        ],
        forest_name="",
        address="",
        output_dir=r"/home/ubuntu/gfcnepal/output",
        crs_epsg="EPSG:32645",
        dem_path=r"/home/ubuntu/gfcnepal/dem/Nepal.tif",
    )

    print("Maps generated successfully.")
    